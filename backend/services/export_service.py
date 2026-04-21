import io
import os
import re
import urllib.request
from docx import Document
from docx.shared import RGBColor
from fpdf import FPDF
from backend.core.export_interface import IExportService

# Brand colors aligned directly with the Frontend CSS theme mapping
SPEAKER_COLORS = [
    (65, 185, 198),   # S0: secondary-green
    (241, 166, 160),  # S1: secondary-pink
    (237, 109, 124),  # S2: primary-coral
    (193, 208, 217),  # S3: secondary-grey
    (2, 104, 115)     # S4: primary-green
]

def format_speaker_label(speaker_raw: str) -> str:
    """Cleans redundant 'Speaker speaker_0' formats into cleanly numbered 'Speaker 0'."""
    speaker_raw = str(speaker_raw)
    # Extract the first integer found in the speaker string
    match = re.search(r'\d+', speaker_raw)
    if match:
        return f"Speaker {int(match.group())}"
    
    # Fallback if no digits are present, just remove duplicate "Speaker" words
    cleaned = re.sub(r'(?i)(speaker\s*)+', 'Speaker ', speaker_raw)
    return cleaned.strip().title()

def merge_diarization(diarization):
    if isinstance(diarization, str):
        return diarization
    
    # Intelligently merge adjacent segments of the same speaker
    merged = []
    sorted_diar = sorted(diarization, key=lambda x: x.get('start', 0.0))
    for seg in sorted_diar:
        # Pre-clean the speaker name before doing any logic
        clean_speaker = format_speaker_label(seg.get('speaker', 'Unknown'))
        
        if merged and merged[-1].get('speaker') == clean_speaker:
            merged[-1]['end'] = seg.get('end', 0.0)
            merged[-1]['text'] += " " + seg.get('text', '').strip()
        else:
            new_seg = seg.copy()
            new_seg['speaker'] = clean_speaker
            merged.append(new_seg)
    return merged

class ExportService(IExportService):
    def __init__(self):
        # Download a Unicode font (Roboto) to support Multilingual chars in PDF
        self.font_path = "Roboto-Regular.ttf"
        if not os.path.exists(self.font_path):
            try:
                url = "https://raw.githubusercontent.com/googlefonts/roboto/main/src/hinted/Roboto-Regular.ttf"
                urllib.request.urlretrieve(url, self.font_path)
            except Exception as e:
                print(f"Warning: Failed to download Unicode font. {e}")

    def export_to_txt(self, data) -> bytes:
        merged = merge_diarization(data)
        if isinstance(merged, str):
            return merged.encode("utf-8")

        lines = []
        for seg in merged:
            speaker = seg.get('speaker', 'Unknown')
            start = seg.get('start', 0.0)
            end = seg.get('end', 0.0)
            lines.append(f"[{start:.2f}s - {end:.2f}s] {speaker}:\n{seg.get('text', '')}")

        return "\n\n".join(lines).encode("utf-8")

    def export_to_docx(self, data) -> bytes:
        doc = Document()
        doc.add_heading('Transcription Results', 0)

        merged = merge_diarization(data)
        if isinstance(merged, str):
            doc.add_paragraph(merged)
        else:
            speaker_map = {}
            for seg in merged:
                speaker = seg.get("speaker", "Unknown")
                if speaker not in speaker_map:
                    speaker_map[speaker] = len(speaker_map)
                
                color_idx = speaker_map[speaker] % len(SPEAKER_COLORS)
                r, g, b = SPEAKER_COLORS[color_idx]

                p = doc.add_paragraph()
                
                # Add Speaker metadata cleanly colored and bolded
                speaker_run = p.add_run(f"[{seg.get('start', 0.0):.2f}s - {seg.get('end', 0.0):.2f}s] {speaker}:\n")
                speaker_run.bold = True
                speaker_run.font.color.rgb = RGBColor(r, g, b)
                
                # Add actual transcription text underneath
                p.add_run(seg.get("text", ""))

        output_fp = io.BytesIO()
        doc.save(output_fp)
        return output_fp.getvalue()

    def export_to_pdf(self, data) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        
        if os.path.exists(self.font_path):
            pdf.add_font("Roboto", "", self.font_path, uni=True)
            pdf.set_font("Roboto", size=11)
        else:
            pdf.set_font("Helvetica", size=11)

        merged = merge_diarization(data)
        if isinstance(merged, str):
            text = merged
            if not os.path.exists(self.font_path):
                text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, text=text)
        else:
            speaker_map = {}
            for seg in merged:
                speaker = seg.get("speaker", "Unknown")
                if speaker not in speaker_map:
                    speaker_map[speaker] = len(speaker_map)
                
                color_idx = speaker_map[speaker] % len(SPEAKER_COLORS)
                r, g, b = SPEAKER_COLORS[color_idx]

                text = seg.get("text", "")
                if not os.path.exists(self.font_path):
                    text = text.encode('latin-1', 'replace').decode('latin-1')
                    speaker_safe = speaker.encode('latin-1', 'replace').decode('latin-1')
                else:
                    speaker_safe = speaker

                # 1. Set Custom Color for Speaker Header
                pdf.set_text_color(r, g, b)
                if not os.path.exists(self.font_path):
                    pdf.set_font("Helvetica", "B", 11)
                    
                pdf.cell(0, 6, txt=f"[{seg.get('start', 0.0):.2f}s - {seg.get('end', 0.0):.2f}s] {speaker_safe}:", ln=True)

                # 2. Reset to Solid Black for the standard Text Body
                pdf.set_text_color(0, 0, 0)
                if not os.path.exists(self.font_path):
                    pdf.set_font("Helvetica", "", 11)
                    
                pdf.multi_cell(0, 6, text=text)
                pdf.ln(4)

        return pdf.output()